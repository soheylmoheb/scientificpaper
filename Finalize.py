#!/usr/bin/env python3
"""
Flexible Word Report Generator – auto‑detects folder structure, parallel Mendeley lookups,
and generates a professional .docx with all analyses.
"""

import os
import sys
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm  # for progress bar

# =============================================================================
# CONFIGURATION – SET ENVIRONMENT VARIABLES (DO NOT HARDCODE)
# =============================================================================
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
MENDELEY_TOKEN = os.environ.get("MENDELEY_TOKEN")

if not DEEPSEEK_API_KEY or not MENDELEY_TOKEN:
    sys.exit("ERROR: Set DEEPSEEK_API_KEY and MENDELEY_TOKEN environment variables.")

DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# =============================================================================
# STEP 1: ASK FOR ROOT FOLDER AND DETECT STRUCTURE
# =============================================================================
def get_root_folder():
    default = "."
    folder = input(f"Enter the path to the root folder (press Enter for current '{default}'): ").strip()
    if not folder:
        folder = default
    if not os.path.isdir(folder):
        sys.exit(f"Folder '{folder}' does not exist.")
    return folder

def detect_structure(root):
    """
    Detect whether the root contains category folders or direct paper folders.
    Returns a list of paper dictionaries with 'category' (None if flat) and 'path'.
    """
    # Expected category folder names
    category_names = [
        "internet_pricing_1990_2000",
        "bandwidth_pricing_1990_2000",
        "internet_pricing_2000_2010",
        "bandwidth_pricing_2000_2010"
    ]
    category_display = {
        "internet_pricing_1990_2000": "internet pricing (from 1990 to 2000)",
        "bandwidth_pricing_1990_2000": "bandwidth pricing (from 1990 to 2000)",
        "internet_pricing_2000_2010": "internet pricing (from 2000 to 2010)",
        "bandwidth_pricing_2000_2010": "bandwidth pricing (from 2000 to 2010)"
    }

    papers = []
    # Check if any category folder exists
    found_categories = [d for d in os.listdir(root) if os.path.isdir(os.path.join(root, d)) and d in category_names]
    if found_categories:
        # Hierarchical structure
        print("Detected category folders.")
        for cat_dir in found_categories:
            cat_path = os.path.join(root, cat_dir)
            cat_display = category_display.get(cat_dir, cat_dir)
            for paper_folder in os.listdir(cat_path):
                paper_path = os.path.join(cat_path, paper_folder)
                if os.path.isdir(paper_path):
                    papers.append({
                        "category": cat_display,
                        "title": paper_folder,
                        "path": paper_path
                    })
    else:
        # Assume every subdirectory is a paper folder (flat structure)
        print("No category folders found. Assuming each subdirectory is a paper.")
        for item in os.listdir(root):
            item_path = os.path.join(root, item)
            if os.path.isdir(item_path):
                papers.append({
                    "category": None,  # no category heading
                    "title": item,
                    "path": item_path
                })
    return papers

# =============================================================================
# STEP 2: READ DEMAND FILES FOR A PAPER
# =============================================================================
def read_demand_files(paper_path):
    """Return dict {1: content, ..., 8: content}."""
    demands = {}
    for i in range(1, 9):
        fname = f"demand_{i:02d}.txt"
        fpath = os.path.join(paper_path, fname)
        if os.path.isfile(fpath):
            with open(fpath, 'r', encoding='utf-8') as f:
                demands[i] = f.read()
        else:
            demands[i] = f"*Missing demand {i}*"
    return demands

# =============================================================================
# STEP 3: MENDELEY METADATA LOOKUP (PARALLEL)
# =============================================================================
def search_mendeley(title):
    """Search Mendeley by title and return metadata dict or None."""
    headers = {"Authorization": f"Bearer {MENDELEY_TOKEN}"}
    params = {"title": title, "limit": 1}
    url = "https://api.mendeley.com/search/documents"
    try:
        resp = requests.get(url, headers=headers, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        if data:
            doc = data[0]
            authors = []
            if 'authors' in doc:
                for a in doc['authors']:
                    if 'first_name' in a and 'last_name' in a:
                        authors.append(f"{a['first_name']} {a['last_name']}")
                    elif 'last_name' in a:
                        authors.append(a['last_name'])
            return {
                "title": doc.get('title', title),
                "authors": authors,
                "year": doc.get('year', ''),
                "journal": doc.get('source', '') or doc.get('journal', '') or ''
            }
    except Exception as e:
        # Silently fail; will use fallback
        pass
    return None

def fetch_all_metadata(papers):
    """Fetch metadata for all papers in parallel."""
    print("\n🔍 Fetching metadata from Mendeley in parallel...")
    with ThreadPoolExecutor(max_workers=5) as executor:
        future_to_paper = {executor.submit(search_mendeley, p['title']): p for p in papers}
        for future in tqdm(as_completed(future_to_paper), total=len(papers), desc="Mendeley lookups"):
            paper = future_to_paper[future]
            try:
                meta = future.result()
                paper['meta'] = meta
            except Exception:
                paper['meta'] = None
    return papers

# =============================================================================
# STEP 4: GENERATE AI SECTIONS (INTRO, DISCUSSION, CONCLUSION)
# =============================================================================
def generate_ai_section(all_text, section_name, num_paragraphs=None, papers_count=0):
    """Use DeepSeek to generate a section based on all papers."""
    prompt = f"""
You are an expert researcher synthesizing findings from multiple scientific papers about internet and bandwidth pricing models (1990–2010).

Below is the combined analysis of {papers_count} papers, each containing 8 detailed demands (model explanation, formulas, algorithms, code, AI suggestions, datasets, and key findings).

Please write a {section_name} for a comprehensive report. """
    if section_name == "introduction":
        prompt += f"Write a {num_paragraphs}-paragraph introduction that sets the context, outlines the importance of pricing models, and previews the content of the report."
    elif section_name == "discussion":
        prompt += "Write a discussion section that compares and contrasts the different models, highlights common themes, methodological differences, and implications. Identify any controversies or gaps."
    elif section_name == "conclusion":
        prompt += f"Write a {num_paragraphs}-paragraph conclusion that summarizes the main insights, suggests future research directions, and reflects on the evolution of pricing models over the two decades."

    prompt += f"\n\nHere is the aggregated content from all papers:\n\n{all_text[:150000]}"  # Truncate for token limits

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "You are a research assistant specializing in economic models of internet pricing."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.4,
        "max_tokens": 4000
    }
    try:
        resp = requests.post(DEEPSEEK_API_URL, json=payload, headers=headers, timeout=60)
        resp.raise_for_status()
        result = resp.json()
        return result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"\n⚠️ Error generating {section_name}: {e}")
        return f"[{section_name} could not be generated due to API error.]"

# =============================================================================
# STEP 5: WORD DOCUMENT BUILDING
# =============================================================================
def set_heading(paragraph, level=1):
    paragraph.style = f'Heading {level}'

def add_code_block(doc, code_text):
    """Insert code block with monospace font and light gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), "F0F0F0")
    p._element.get_or_add_pPr().append(shading_elm)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def build_document(papers, all_text, output_filename="DeepSeek_Analysis_Report.docx"):
    print("\n📄 Building Word document...")
    doc = Document()

    # Subsections mapping
    subsections = {
        1: "a) Explanation of the pricing model proposed in this paper",
        2: "b) Mathematical formulas",
        3: "c) Step‑by‑step algorithm to perform a Monte Carlo simulation",
        4: "d) Python code that implements the Monte Carlo simulation",
        5: "e) Best machine learning or AI algorithm to predict internet prices based on this model",
        6: "f) Python code for that AI algorithm",
        7: "g) Kaggle dataset (or library) that could be used to train the AI model",
        8: "h) Key findings"
    }

    # Collect citations
    citations = []

    # ---- INTRODUCTION (will be generated later, insert now) ----
    print("  Generating introduction...")
    intro = generate_ai_section(all_text, "introduction", 5, len(papers))
    p = doc.add_paragraph("Introduction", style='Heading 1')
    doc.add_paragraph(intro)

    # ---- MAIN BODY ----
    # Group papers by category (if any)
    categories = {}
    for p in papers:
        cat = p.get('category', 'Uncategorized')
        if cat is None:
            cat = 'All Papers'
        categories.setdefault(cat, []).append(p)

    for cat_name, cat_papers in categories.items():
        if cat_name != 'All Papers':
            doc.add_paragraph(cat_name, style='Heading 1')
        for paper in cat_papers:
            # Citation string
            meta = paper.get('meta')
            if meta and meta.get('authors'):
                authors = ", ".join(meta['authors'])
                year = meta.get('year', '')
                title = meta.get('title', paper['title'])
                journal = meta.get('journal', '')
                cit = f"{authors} ({year}). {title}. {journal}."
            else:
                cit = f"{paper['title']} (metadata not found)."
            citations.append(cit)

            # Paper title as Heading 2 with citation number
            doc.add_paragraph(f"{paper['title']} [{len(citations)}]", style='Heading 2')

            # Read demand files
            demands = read_demand_files(paper['path'])

            # Subsections
            for i in range(1, 9):
                doc.add_paragraph(subsections[i], style='Heading 3')
                content = demands[i]
                if i in [4, 6]:  # code demands
                    add_code_block(doc, content)
                else:
                    doc.add_paragraph(content)

    # ---- DISCUSSION ----
    print("  Generating discussion...")
    discussion = generate_ai_section(all_text, "discussion", None, len(papers))
    doc.add_paragraph("Discussion", style='Heading 1')
    doc.add_paragraph(discussion)

    # ---- CONCLUSION ----
    print("  Generating conclusion...")
    conclusion = generate_ai_section(all_text, "conclusion", 3, len(papers))
    doc.add_paragraph("Conclusion", style='Heading 1')
    doc.add_paragraph(conclusion)

    # ---- BIBLIOGRAPHY ----
    doc.add_paragraph("Bibliography", style='Heading 1')
    for idx, cit in enumerate(citations, 1):
        doc.add_paragraph(f"[{idx}] {cit}")

    # Save
    doc.save(output_filename)
    print(f"\n✅ Report saved as '{output_filename}'")

# =============================================================================
# MAIN
# =============================================================================
def main():
    print("="*70)
    print("FLEXIBLE WORD REPORT GENERATOR – AUTO‑STRUCTURE DETECTION")
    print("="*70)

    root = get_root_folder()
    papers = detect_structure(root)
    if not papers:
        sys.exit("No paper folders found. Please ensure your folder contains subdirectories with demand_*.txt files.")

    print(f"Found {len(papers)} paper folders.")

    # Read all text for AI sections (combine all demand files)
    all_text = ""
    for p in papers:
        demands = read_demand_files(p['path'])
        all_text += f"\n\n--- PAPER: {p['title']} ---\n"
        for i in range(1, 9):
            all_text += f"\nDemand {i}:\n{demands[i]}\n"

    # Fetch Mendeley metadata in parallel
    papers = fetch_all_metadata(papers)

    # Build document
    build_document(papers, all_text)

if __name__ == "__main__":
    main()
